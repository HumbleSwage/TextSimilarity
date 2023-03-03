from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator
from docx import Document


class readFileContent:
    """
    @todo:将pdf文件处理成要求格式；将word文件处理成要求格式；将txt文件处理成要求格式
    @return:content list 文件内容 [文件内容]
    @param:_pdfPath指定pdf文件的绝对路径
           _wordWord指定word文件的绝对路径
           _txtPath指定txt文件的绝对路径
    """

    def __init__(self):
        self.content = ""
        self.msg = 1

    def readPdf(self, Path):
        try:
            Path = open(Path, 'rb')
            self.content = ""
            # 来创建一个pdf文档分析器
            parser = PDFParser(Path)
            # 创建一个PDF文档对象存储文档结构
            document = PDFDocument(parser)
            # 检查文件是否允许文本提取
            if not document.is_extractable:
                raise PDFTextExtractionNotAllowed
            else:
                # 创建一个PDF资源管理器对象来存储共赏资源
                rsrcmgr = PDFResourceManager()
                # 设定参数进行分析
                laparams = LAParams()
                # 创建一个PDF设备对象
                device = PDFPageAggregator(rsrcmgr, laparams=laparams)
                # 创建一个PDF解释器对象
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                # 处理每一页
                for page in PDFPage.create_pages(document):
                    strEachPage = ""
                    interpreter.process_page(page)
                    # 接受该页面的LTPage对象
                    layout = device.get_result()
                    for x in layout:
                        if isinstance(x, LTTextBoxHorizontal):
                            strEachPage = strEachPage + x.get_text()
                    self.content = self.content + strEachPage
        except UnicodeDecodeError:
            print(1)
            self.msg = 0
        finally:
            return self.content, self.msg

    def readWord(self, _wordWord):
        try:
            self.content = ""
            d = Document(_wordWord)
            # 遍历段落
            for paragraph in d.paragraphs:
                self.content = self.content.strip() + paragraph.text
            # 遍历表格
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.content = self.content.strip() + cell.text
        except UnicodeDecodeError:
            self.msg = 0
        finally:
            return self.content, self.msg

    def readTxt(self, _txtPath):
        try:
            self.content = ""
            with open(_txtPath, mode="r") as txt:
                for i in txt.readlines():
                    self.content = self.content + i
        except UnicodeDecodeError:
            self.msg = 0
        finally:
            return self.content, self.msg
